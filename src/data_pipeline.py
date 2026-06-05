import os
import re
import json
import logging
import pandas as pd
import numpy as np

# Set up module-level logger
logger = logging.getLogger(__name__)

def docx_to_text(file_path: str) -> str:
    """
    Extracts raw text strings from a job description DOCX file using python-docx.
    
    Parameters:
        file_path (str): Path to the job description DOCX file.
        
    Returns:
        str: Consolidated raw text from the document.
    """
    if not os.path.exists(file_path):
        logger.warning(f"DOCX file not found at: {file_path}")
        return ""
    
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        # Include tables text if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join([p.strip() for p in paragraphs if p.strip()])
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {e}")
        return ""

class TalentDataPipeline:
    """
    Data ingestion, cleaning, and feature engineering pipeline for candidates.
    Supports legacy CSV formats as well as nested JSON candidate profiles and DOCX jobs.
    """
    def __init__(self):
        self.raw_profiles = None
        self.raw_jobs = None
        self.clean_profiles = None
        self.clean_jobs = None

    def load_data(self, profiles_path: str = 'data/raw_profiles.csv', jobs_path: str = 'data/job_descriptions.csv') -> tuple:
        """
        Loads candidate profiles and job descriptions from CSV files (legacy).
        """
        if not os.path.exists(profiles_path):
            raise FileNotFoundError(f"Candidate profiles file not found at: {os.path.abspath(profiles_path)}")
        if not os.path.exists(jobs_path):
            raise FileNotFoundError(f"Job descriptions file not found at: {os.path.abspath(jobs_path)}")
            
        try:
            self.raw_profiles = pd.read_csv(profiles_path)
        except Exception as e:
            raise IOError(f"Error loading candidate profiles CSV from {profiles_path}: {e}")
            
        try:
            self.raw_jobs = pd.read_csv(jobs_path)
        except Exception as e:
            raise IOError(f"Error loading job descriptions CSV from {jobs_path}: {e}")
            
        return self.raw_profiles, self.raw_jobs

    def handle_missing_values(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Handles missing values in a DataFrame:
        - Imputes numeric fields with their median values.
        - Fills categorical/text fields (non-numeric) with empty strings.
        """
        if df is None:
            return None
            
        df_clean = df.copy()
        
        numeric_cols = df_clean.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            median_val = df_clean[col].median()
            if pd.isna(median_val):
                median_val = 0.0
            df_clean[col] = df_clean[col].fillna(median_val)
            
        categorical_cols = df_clean.select_dtypes(exclude=[np.number]).columns
        for col in categorical_cols:
            df_clean[col] = df_clean[col].fillna("")
            
        return df_clean

    def normalize_text(self, text: str) -> str:
        """
        Normalizes experience and description text:
        - Converts all characters to lowercase.
        - Removes special regex markers and symbols.
        - Standardizes whitespaces and trims margins.
        """
        if not isinstance(text, str):
            return ""
        # Lowercase
        normalized = text.lower()
        # Remove special regex markers and symbols
        normalized = re.sub(r'[^\w\s-]', ' ', normalized)
        # Standardize whitespace
        normalized = re.sub(r'\s+', ' ', normalized)
        return normalized.strip()

    def normalize_experience_fields(self, df: pd.DataFrame, columns: list = None) -> pd.DataFrame:
        """
        Normalizes experience text fields in the DataFrame.
        """
        if df is None:
            return None
            
        df_clean = df.copy()
        
        if columns is None:
            columns = [col for col in df_clean.columns if any(k in col.lower() for k in ['experience', 'description', 'profile', 'text']) and not pd.api.types.is_numeric_dtype(df_clean[col])]
            
        for col in columns:
            if col in df_clean.columns:
                df_clean[col] = df_clean[col].astype(str).apply(self.normalize_text)
                
        return df_clean

    # -------------------------------------------------------------------------
    # PROXY-SKILL TRANSLATION MATRIX
    # Maps non-standard technical vernacular from Tier-2/Tier-3 Indian regional
    # college projects to global industry-standard equivalents.
    # Sorted longest-first at runtime to avoid substring collision.
    # -------------------------------------------------------------------------
    PROXY_SKILL_MAP: dict = {
        # --- Campus / Infrastructure ---
        "completed institute campus networking layout": "Network Architecture",
        "institute campus networking layout": "Network Architecture",
        "campus networking layout": "Network Architecture",
        "local institute network": "Network Architecture",
        "college lan": "Network Architecture",
        "campus network": "Network Architecture",
        "lan setup": "Network Architecture",
        "cyber café management": "IT Support",
        # --- Database & Back-end ---
        "handled local store ledger system database": "Database Administration",
        "local store ledger system database": "Database Administration",
        "ledger system database": "Database Administration",
        "local store inventory system": "Database Administration",
        "ms access database project": "Database Administration",
        "foxpro database": "Database Administration",
        # --- Web / Full-Stack ---
        "designed college web portal via php/mysql": "Full-Stack Web Development",
        "college web portal via php/mysql": "Full-Stack Web Development",
        "college web portal php": "Full-Stack Web Development",
        "php mysql": "Full-Stack Web Development",
        "codeigniter": "PHP Web Framework",
        "dot net": ".NET Development",
        "dot-net": ".NET Development",
        "asp net": ".NET Development",
        "vb project": "Visual Basic Development",
        "vb.net": ".NET Development",
        "j2ee": "Java Enterprise Development",
        "core java": "Java Development",
        "struts framework": "Java Web Framework",
        "spring mvc": "Java Web Framework",
        # --- Networking / Admin ---
        "cisco packet tracer project": "Network Architecture",
        "nsit networking lab": "Network Administration",
        "lan setup": "Network Administration",
        # --- Office / ERP ---
        "tally": "Accounting Software",
        "tally erp": "ERP Systems",
        "busy accounting": "Accounting Software",
        "ms office advanced": "Productivity Suite",
        "computer operator": "Software Operations",
        # --- Data / Analytics ---
        "excel analytics": "Data Analysis",
        "excel macros": "Data Analysis Automation",
        "spss": "Statistical Analysis",
        "r programming": "Statistical Programming",
        "minitab": "Statistical Analysis",
        # --- Hardware / IT Support ---
        "hardware maintenance": "IT Support",
        "assembling pc": "IT Support",
        "printer maintenance": "IT Support",
        "dtp": "Desktop Publishing",
        "dtp operator": "Desktop Publishing",
        # --- Cloud / DevOps ---
        "aws free tier project": "Cloud Computing",
        "digitalocean droplet": "Cloud Computing",
        "heroku deployment": "Cloud Deployment",
        # --- Security ---
        "ethical hacking course": "Cybersecurity",
        "kali linux basics": "Cybersecurity",
        "cyber security certificate": "Cybersecurity",
        # --- General ---
        "data entry": "Data Management",
        "internet browsing": "Digital Literacy",
        "social media management": "Digital Marketing",
        "seo basics": "Search Engine Optimization",
    }

    # Columns containing demographic/identifying attributes to drop
    _PII_COLUMNS: tuple = (
        'name', 'full_name', 'first_name', 'last_name',
        'email', 'email_address', 'phone', 'phone_number', 'mobile',
        'gender', 'sex', 'dob', 'date_of_birth', 'age',
        'address', 'city', 'state', 'country', 'pincode', 'zip_code',
        'photo', 'photo_url', 'linkedin_url', 'github_url',
        'religion', 'caste', 'nationality', 'marital_status',
        'anonymized_name', 'location'
    )

    def load_and_clean_candidates(
        self,
        json_path: str,
        max_rows: int = None,
    ) -> pd.DataFrame:
        """
        Loads nested JSON/JSONL candidates into a flat pandas DataFrame targeting the official schema.

        Memory strategy
        ---------------
        - **JSONL** (`.jsonl`, or any file whose first non-whitespace character is
          not ``[``): records are yielded one line at a time so the full file is
          never held in RAM simultaneously.  Safe for 100k+ candidate pools.
        - **JSON array** (`.json`, first char ``[``): loaded with ``json.load``
          in a single pass — appropriate for smaller sample files (≤ a few MB).

        Parameters
        ----------
        json_path : str
            Path to the candidate data file (.json or .jsonl).
        max_rows : int, optional
            If set, stop reading after this many successfully parsed records.
            Use for trial/debug runs to avoid ingesting the full dataset.
            Default is None (read entire file).

        Handles missing keys and applies median/zero imputation.
        """
        if not os.path.exists(json_path):
            raise FileNotFoundError(
                f"JSON candidate file not found at: {os.path.abspath(json_path)}"
            )

        def _stream_records(path: str):
            """Generator: yields one raw dict per candidate, never buffering the whole file."""
            with open(path, 'r', encoding='utf-8') as fh:
                # Peek at first non-whitespace byte to detect format
                first_char = ''
                while True:
                    ch = fh.read(1)
                    if not ch:
                        return
                    if not ch.isspace():
                        first_char = ch
                        break

                if first_char == '[':
                    # JSON array — read remainder and parse whole document
                    rest = fh.read()
                    try:
                        arr = json.loads(first_char + rest)
                    except json.JSONDecodeError as exc:
                        raise IOError(f"JSON parse error in {path}: {exc}") from exc
                    for item in arr:
                        if isinstance(item, dict):
                            yield item
                else:
                    # JSONL — true line-by-line streaming (RAM-safe for 100k+ records).
                    # first_char is the first non-whitespace byte of line 1;
                    # read the remainder of that line, then iterate the rest lazily.
                    import itertools
                    rest_of_line1 = fh.readline()          # reads up to and including \n
                    line1 = first_char + rest_of_line1
                    for line in itertools.chain([line1], fh):   # lazy: fh yields one line at a time
                        stripped = line.strip()
                        if stripped:
                            try:
                                obj = json.loads(stripped)
                                if isinstance(obj, dict):
                                    yield obj
                            except json.JSONDecodeError:
                                pass  # skip malformed lines

        parsed = []
        for cand in _stream_records(json_path):
            # Early-exit for trial / debug runs — stops reading the file
            # immediately once max_rows records have been collected.
            if max_rows is not None and len(parsed) >= max_rows:
                break

            cand_id = cand.get('candidate_id', '')
            
            # Extract Profile (Nested)
            profile = cand.get('profile')
            if not isinstance(profile, dict):
                profile = {}
            
            anonymized_name = profile.get('anonymized_name') or ''
            headline = profile.get('headline') or ''
            summary = profile.get('summary') or ''
            location = profile.get('location') or ''
            country = profile.get('country') or ''
            current_title = profile.get('current_title') or ''
            current_company = profile.get('current_company') or ''
            current_company_size = profile.get('current_company_size') or ''
            current_industry = profile.get('current_industry') or ''
            
            try:
                yoe_raw = profile.get('years_of_experience')
                years_of_experience = float(yoe_raw) if yoe_raw is not None else np.nan
            except (ValueError, TypeError):
                years_of_experience = np.nan
            
            # Extract Career History (Array of Dicts)
            career_history = cand.get('career_history')
            if not isinstance(career_history, list):
                career_history = []
            
            durations = []
            descriptions = []
            for job in career_history:
                if not isinstance(job, dict):
                    continue
                desc = job.get('description') or ''
                if desc:
                    descriptions.append(desc)
                
                try:
                    dur_raw = job.get('duration_months')
                    if dur_raw is not None:
                        durations.append(float(dur_raw))
                except (ValueError, TypeError):
                    pass
            
            # Compute Trajectory Metrics
            if durations:
                total_career_duration = sum(durations) / 12.0
                avg_tenure_per_role = (sum(durations) / len(durations)) / 12.0
                promotion_velocity = len(durations) / total_career_duration if total_career_duration > 0 else 0.0
            else:
                total_career_duration = 0.0
                avg_tenure_per_role = 0.0
                promotion_velocity = 0.0
                
            # Build capability text blob
            text_parts = [headline, summary, current_title, current_industry] + descriptions
            experience_text = " ".join([str(t).strip() for t in text_parts if t])
            
            # Extract Skills — support both list-of-strings and list-of-dicts (with 'name' key)
            skills = cand.get('skills')
            if isinstance(skills, list):
                skill_names = []
                for s in skills:
                    if isinstance(s, dict):
                        skill_names.append(str(s.get('name', '')).strip())
                    elif s:
                        skill_names.append(str(s).strip())
                skills_raw = ", ".join([n for n in skill_names if n])
            elif isinstance(skills, str):
                skills_raw = skills
            else:
                skills_raw = ""
                
            # Extract Redrob Behavioral Signals (Nested)
            signals = cand.get('redrob_signals')
            if not isinstance(signals, dict):
                signals = {}
                
            try:
                profile_completeness_score = float(signals.get('profile_completeness_score', np.nan))
            except (ValueError, TypeError):
                profile_completeness_score = np.nan
                
            open_to_work = signals.get('open_to_work_flag')
            try:
                if isinstance(open_to_work, str):
                    open_to_work_flag = 1.0 if open_to_work.lower() in ('true', '1') else 0.0
                else:
                    open_to_work_flag = 1.0 if open_to_work else 0.0
            except (ValueError, TypeError):
                open_to_work_flag = 0.0
                
            try:
                profile_views = float(signals.get('profile_views_received_30d', np.nan))
            except (ValueError, TypeError):
                profile_views = np.nan
                
            try:
                apps_submitted = float(signals.get('applications_submitted_30d', np.nan))
            except (ValueError, TypeError):
                apps_submitted = np.nan
                
            try:
                recruiter_resp = float(signals.get('recruiter_response_rate', np.nan))
            except (ValueError, TypeError):
                recruiter_resp = np.nan
                
            try:
                avg_resp_time = float(signals.get('avg_response_time_hours', np.nan))
            except (ValueError, TypeError):
                avg_resp_time = np.nan
                
            try:
                connections = float(signals.get('connection_count', np.nan))
            except (ValueError, TypeError):
                connections = np.nan
                
            assess_scores = signals.get('skill_assessment_scores')
            mean_assessment_score = np.nan
            if isinstance(assess_scores, dict):
                vals = []
                for v in assess_scores.values():
                    try:
                        vals.append(float(v))
                    except (ValueError, TypeError):
                        pass
                if vals:
                    mean_assessment_score = float(np.mean(vals))
            elif isinstance(assess_scores, list):
                vals = []
                for v in assess_scores:
                    try:
                        vals.append(float(v))
                    except (ValueError, TypeError):
                        pass
                if vals:
                    mean_assessment_score = float(np.mean(vals))
            elif assess_scores is not None:
                try:
                    mean_assessment_score = float(assess_scores)
                except (ValueError, TypeError):
                    pass
            
            # Days since last active (ref date: 2026-06-04)
            last_active = signals.get('last_active_date')
            signup = signals.get('signup_date')
            days_since_last_active = np.nan
            if last_active:
                try:
                    dt_active = pd.to_datetime(last_active)
                    dt_ref = pd.to_datetime('2026-06-04')
                    days_since_last_active = (dt_ref - dt_active).total_seconds() / 86400.0
                except Exception:
                    pass
            days_since_signup = np.nan
            if signup:
                try:
                    dt_signup = pd.to_datetime(signup)
                    dt_ref = pd.to_datetime('2026-06-04')
                    days_since_signup = (dt_ref - dt_signup).total_seconds() / 86400.0
                except Exception:
                    pass
            
            # Map downstream behavior metrics
            profile_update_frequency = profile_completeness_score
            interaction_history_index = connections
            submission_timestamp_index = -days_since_last_active if not pd.isna(days_since_last_active) else np.nan
            
            row = {
                'candidate_id': cand_id,
                # PII fields kept here so engineer_features can audit/drop them explicitly
                'anonymized_name': anonymized_name,
                'location': location,
                'country': country,
                # Capability text blobs
                'headline_raw': headline,
                'summary_raw': summary,
                'experience_text': experience_text,
                'skills_raw': skills_raw,
                # Trajectory (duration_months-based, pre-computed for speed)
                'years_of_experience': years_of_experience,
                'total_career_duration': total_career_duration,
                'avg_tenure_per_role': avg_tenure_per_role,
                'promotion_velocity': promotion_velocity,
                # Raw career history stored as JSON string for date-based re-parsing in engineer_features
                'career_history_raw': json.dumps(career_history) if career_history else '[]',
                # Behavioral signals
                'profile_completeness_score': profile_completeness_score,
                'open_to_work_flag': open_to_work_flag,
                'profile_views_received_30d': profile_views,
                'applications_submitted_30d': apps_submitted,
                'recruiter_response_rate': recruiter_resp,
                'avg_response_time_hours': avg_resp_time,
                'connection_count': connections,
                'skill_assessment_scores': mean_assessment_score,
                'days_since_last_active': days_since_last_active,
                'days_since_signup': days_since_signup,
                'profile_update_frequency': profile_update_frequency,
                'submission_timestamp_index': submission_timestamp_index,
                'interaction_history_index': interaction_history_index
            }
            parsed.append(row)
            
        df = pd.DataFrame(parsed)
        
        # Impute missing values
        if not df.empty:
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            for col in numeric_cols:
                median_val = df[col].median()
                if pd.isna(median_val):
                    median_val = 0.0
                df[col] = df[col].fillna(median_val)
                
            text_cols = df.select_dtypes(exclude=[np.number]).columns
            for col in text_cols:
                df[col] = df[col].fillna("")
        else:
            # Handle empty result schema
            df = pd.DataFrame(columns=[
                'candidate_id', 'anonymized_name', 'location', 'country',
                'headline_raw', 'summary_raw', 'experience_text', 'skills_raw',
                'years_of_experience', 'total_career_duration', 'avg_tenure_per_role',
                'promotion_velocity', 'career_history_raw',
                'profile_completeness_score', 'open_to_work_flag',
                'profile_views_received_30d', 'applications_submitted_30d',
                'recruiter_response_rate', 'avg_response_time_hours', 'connection_count',
                'skill_assessment_scores', 'days_since_last_active', 'days_since_signup',
                'profile_update_frequency', 'submission_timestamp_index', 'interaction_history_index'
            ])
            
        df = self.sweep_honeypots(df)
        return df

    def sweep_honeypots(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Sweeps out honeypot candidate profiles that represent fake, dummy, or test users.
        Filters out rows where:
        - candidate_id or anonymized_name contains 'honeypot', 'fake', 'dummy', 'test_user', or 'spam' (case-insensitive).
        - years_of_experience is negative, null, or greater than 60.
        - the profile has no skills and no experience text.
        """
        if df is None or df.empty:
            return df
        
        initial_count = len(df)
        
        # 1. Check ID and Name for test/fake/honeypot keywords
        id_mask = df['candidate_id'].astype(str).str.lower().str.contains('honeypot|fake|dummy|spam|test_user', regex=True)
        name_mask = pd.Series(False, index=df.index)
        if 'anonymized_name' in df.columns:
            name_mask = df['anonymized_name'].astype(str).str.lower().str.contains('honeypot|fake|dummy|spam|test_user', regex=True)
            
        # 2. Check for extreme/invalid experience values
        exp_mask = pd.Series(False, index=df.index)
        if 'years_of_experience' in df.columns:
            exp_values = pd.to_numeric(df['years_of_experience'], errors='coerce')
            exp_mask = (exp_values < 0) | (exp_values > 60) | exp_values.isna()
            
        # 3. Check for completely empty profiles (no skills and no experience text)
        empty_mask = pd.Series(False, index=df.index)
        skills_col = 'skills_raw' if 'skills_raw' in df.columns else ('skills_clean' if 'skills_clean' in df.columns else None)
        exp_text_col = 'experience_text' if 'experience_text' in df.columns else None
        
        if skills_col and exp_text_col:
            empty_mask = (df[skills_col].astype(str).str.strip() == '') & (df[exp_text_col].astype(str).str.strip() == '')
            
        # Combine masks for removal
        honeypot_mask = id_mask | name_mask | exp_mask | empty_mask
        
        cleaned_df = df[~honeypot_mask].copy()
        removed_count = initial_count - len(cleaned_df)
        if removed_count > 0:
            logger.info(f"Swept {removed_count} honeypots out of the dataset.")
            
        return cleaned_df

    # PII column set extended with raw field names introduced in load_and_clean_candidates
    _PII_COLUMNS: tuple = (
        'name', 'full_name', 'first_name', 'last_name',
        'email', 'email_address', 'phone', 'phone_number', 'mobile',
        'gender', 'sex', 'dob', 'date_of_birth', 'age',
        'address', 'city', 'state', 'country', 'pincode', 'zip_code',
        'photo', 'photo_url', 'linkedin_url', 'github_url',
        'religion', 'caste', 'nationality', 'marital_status',
        'anonymized_name', 'location',
        # Raw text staging cols that carry location/person signals
        'headline_raw', 'summary_raw', 'career_history_raw',
    )

    # Reference date used for computing recency from date fields
    _REF_DATE: pd.Timestamp = pd.Timestamp('2026-06-04')

    def engineer_features(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Comprehensive feature engineering stage.

        1. **Career Trajectory & Velocity**
           Parses `career_history_raw` (JSON list of dicts) using ISO `start_date` /
           `end_date` fields to derive date-accurate duration metrics:
             - `total_duration_months`  : cumulative employment in calendar months
             - `num_transitions`        : number of unique job transitions recorded
             - `avg_tenure_months`      : mean months per role
             - `career_velocity_score`  : transitions / (total_duration_months / 12),
                                         zero-safe; captures promotion cadence.

        2. **Proxy-Skill Translation Matrix**
           Applies PROXY_SKILL_MAP over three text fields
           ('skills_raw', 'headline_raw', 'summary_raw') to normalise
           Tier-2/Tier-3 Indian vernacular into global industry standards,
           writing results to 'skills_clean' and 'text_normalized'.

        3. **Anonymization Guardrails**
           Drops all PII / demographic columns (names, location, country,
           raw staging fields) so downstream ranking evaluates candidates
           strictly on merit and behavioral signals.

        Parameters
        ----------
        df : pd.DataFrame
            Output of `load_and_clean_candidates` (or legacy CSV pipeline).

        Returns
        -------
        pd.DataFrame
            Anonymised, feature-engineered DataFrame ready for vectorisation.
        """
        if df is None or df.empty:
            return df

        feat = df.copy()

        # =================================================================
        # 1. DATE-BASED CAREER TRAJECTORY & VELOCITY
        # =================================================================
        _REF = self._REF_DATE

        def _parse_date(val: str) -> pd.Timestamp:
            """Parse partial ISO dates ('YYYY-MM', 'YYYY-MM-DD') robustly."""
            try:
                return pd.to_datetime(str(val), format='%Y-%m', errors='raise')
            except Exception:
                try:
                    return pd.to_datetime(str(val), infer_datetime_format=True, errors='coerce')
                except Exception:
                    return pd.NaT

        def _extract_trajectory(raw_json: str) -> dict:
            """
            Parse one candidate's career_history JSON string and return
            a dict of computed trajectory scalars.
            """
            empty = {
                'total_duration_months': 0.0,
                'num_transitions': 0,
                'avg_tenure_months': 0.0,
                'career_velocity_score': 0.0,
            }
            try:
                history = json.loads(raw_json) if isinstance(raw_json, str) else raw_json
            except (json.JSONDecodeError, TypeError):
                return empty

            if not isinstance(history, list) or len(history) == 0:
                return empty

            role_months: list = []
            for job in history:
                if not isinstance(job, dict):
                    continue

                is_current = bool(job.get('is_current', False))
                start_raw  = job.get('start_date')
                end_raw    = job.get('end_date', '')

                # Prefer duration_months when present (exact value from data)
                dur_raw = job.get('duration_months')
                if dur_raw is not None:
                    try:
                        role_months.append(float(dur_raw))
                        continue
                    except (ValueError, TypeError):
                        pass

                # Fall back to date subtraction
                t_start = _parse_date(start_raw)
                if pd.isna(t_start):
                    continue

                if is_current or str(end_raw).strip().lower() in ('present', 'current', '', 'none', 'null'):
                    t_end = _REF
                else:
                    t_end = _parse_date(end_raw)
                    if pd.isna(t_end):
                        t_end = _REF

                # Month delta = year diff * 12 + month diff
                months = max(
                    (t_end.year - t_start.year) * 12 + (t_end.month - t_start.month),
                    0,
                )
                role_months.append(float(months))

            if not role_months:
                return empty

            total_months = float(np.sum(role_months))
            n_transitions = len(role_months)
            avg_months    = float(np.mean(role_months))
            total_years   = total_months / 12.0
            velocity      = round(n_transitions / total_years, 4) if total_years > 0 else 0.0

            return {
                'total_duration_months': total_months,
                'num_transitions': n_transitions,
                'avg_tenure_months': avg_months,
                'career_velocity_score': velocity,
            }

        if 'career_history_raw' in feat.columns:
            traj_series = feat['career_history_raw'].fillna('[]').apply(_extract_trajectory)
            traj_df = pd.DataFrame(traj_series.tolist(), index=feat.index)
            # Assign — overwrite precomputed cols for date-accuracy; keep old ones as backup
            feat['total_duration_months']   = traj_df['total_duration_months']
            feat['num_transitions']         = traj_df['num_transitions'].astype(int)
            feat['avg_tenure_months']       = traj_df['avg_tenure_months']
            feat['career_velocity_score']   = traj_df['career_velocity_score']
            # Align legacy column names used by downstream modules
            feat['total_career_duration']   = feat['total_duration_months'] / 12.0
            feat['avg_tenure_per_role']     = feat['avg_tenure_months'] / 12.0
            feat['promotion_velocity']      = feat['career_velocity_score']
        else:
            # No raw JSON — use whatever trajectory cols are already present
            for col, default in [
                ('total_duration_months', 0.0), ('num_transitions', 0),
                ('avg_tenure_months', 0.0), ('career_velocity_score', 0.0),
            ]:
                if col not in feat.columns:
                    feat[col] = default

        # =================================================================
        # 2. PROXY-SKILL TRANSLATION (multi-field: skills + headline + summary)
        # =================================================================
        # Pre-sort map longest-phrase-first to avoid substring false-positives
        skill_map_sorted = dict(
            sorted(self.PROXY_SKILL_MAP.items(), key=lambda kv: len(kv[0]), reverse=True)
        )
        # Compile a single multi-pattern regex for speed
        _pattern_cache: list = [
            (re.compile(r'(?i)' + r'\b' + re.escape(src) + r'\b'), tgt)
            for src, tgt in skill_map_sorted.items()
        ]

        def _apply_proxy_map(text: str) -> str:
            """Run all proxy substitutions on a single string (case-insensitive)."""
            if not isinstance(text, str) or not text.strip():
                return ""
            result = text
            for pattern, replacement in _pattern_cache:
                result = pattern.sub(replacement, result)
            return re.sub(r'\s+', ' ', result).strip()

        # skills_clean — translated skills field for downstream embedding
        skills_col = 'skills_raw' if 'skills_raw' in feat.columns else None
        if skills_col:
            feat['skills_clean'] = feat[skills_col].astype(str).apply(_apply_proxy_map)
        else:
            feat['skills_clean'] = ""

        # text_normalized — merged, proxy-translated capability blob for the embedder
        # Combines translated skills + original experience text
        text_components: list = []
        if 'experience_text' in feat.columns:
            text_components.append(feat['experience_text'].astype(str))
        if 'skills_clean' in feat.columns:
            text_components.append(feat['skills_clean'].astype(str))
        # Translate any lingering vernacular inside headline/summary before merging
        for src_col in ('headline_raw', 'summary_raw'):
            if src_col in feat.columns:
                text_components.append(feat[src_col].astype(str).apply(_apply_proxy_map))

        if text_components:
            feat['text_normalized'] = (
                pd.concat(text_components, axis=1)
                .apply(lambda row: ' '.join(v for v in row if v.strip()), axis=1)
            )
        else:
            feat['text_normalized'] = ""

        # =================================================================
        # 3. ANONYMIZATION GUARDRAILS
        #    Drop PII / demographic columns including raw staging fields.
        #    career_history_raw removed here after trajectory extraction.
        # =================================================================
        _pii_set = {p.lower() for p in self._PII_COLUMNS}
        pii_present = [col for col in feat.columns if col.lower().strip() in _pii_set]
        if pii_present:
            feat.drop(columns=pii_present, inplace=True)
            logger.info("PII columns dropped: %s", pii_present)

        return feat

    def run(self, profiles_path: str = 'data/raw_profiles.csv', jobs_path: str = 'data/job_descriptions.csv') -> tuple:
        """
        Master method to execute the data pipeline.
        Supports both CSV schemas and JSON candidate profiles / DOCX job descriptions.
        """
        # Determine format for candidates
        if profiles_path.lower().endswith(('.json', '.jsonl')):
            self.clean_profiles = self.load_and_clean_candidates(profiles_path)
        else:
            raw_profiles, _ = self.load_data(profiles_path, jobs_path)
            self.clean_profiles = self.handle_missing_values(raw_profiles)
            
        # Determine format for job descriptions
        if jobs_path.lower().endswith('.docx'):
            job_text = docx_to_text(jobs_path)
            self.clean_jobs = pd.DataFrame({'job_description': [job_text]})
        else:
            _, raw_jobs = self.load_data(profiles_path, jobs_path)
            self.clean_jobs = self.handle_missing_values(raw_jobs)
            
        # Normalize text/experience fields
        self.clean_profiles = self.normalize_experience_fields(self.clean_profiles)
        self.clean_jobs = self.normalize_experience_fields(self.clean_jobs)
        
        return self.clean_profiles, self.clean_jobs

if __name__ == "__main__":
    import docx
    
    # Setup test paths
    os.makedirs('data', exist_ok=True)
    json_test_file = 'data/sample_candidates.json'
    docx_test_file = 'data/job_description.docx'
    
    # Generate mock JSON if not present
    if not os.path.exists(json_test_file):
        print(f"Generating synthetic candidate profiles in '{json_test_file}'...")
        synthetic_candidates = [
            {
                "candidate_id": "C-1001",
                "profile": {
                    "anonymized_name": "Candidate A",
                    "headline": "Senior Python Engineer",
                    "summary": "Experienced Python developer specializing in REST APIs and cloud architecture.",
                    "location": "Bangalore",
                    "country": "India",
                    "years_of_experience": 6.5,
                    "current_title": "Senior Engineer",
                    "current_company": "Tech Innovations",
                    "current_company_size": "50-200",
                    "current_industry": "Software"
                },
                "career_history": [
                    {
                        "company": "Tech Innovations",
                        "title": "Senior Engineer",
                        "start_date": "2023-01",
                        "end_date": "present",
                        "duration_months": 41,
                        "is_current": True,
                        "industry": "Software",
                        "company_size": "50-200",
                        "description": "Developed serverless APIs using AWS Lambda and API Gateway."
                    },
                    {
                        "company": "Core Systems",
                        "title": "Software Developer",
                        "start_date": "2020-06",
                        "end_date": "2022-12",
                        "duration_months": 30,
                        "is_current": False,
                        "industry": "IT Services",
                        "company_size": "1000+",
                        "description": "Built microservices in Python. Deployed with Docker."
                    }
                ],
                "skills": ["Python", "AWS", "REST APIs", "Docker", "local institute network"],
                "redrob_signals": {
                    "profile_completeness_score": 0.95,
                    "signup_date": "2020-05-15",
                    "last_active_date": "2026-06-02",
                    "open_to_work_flag": True,
                    "profile_views_received_30d": 45,
                    "applications_submitted_30d": 8,
                    "recruiter_response_rate": 0.85,
                    "avg_response_time_hours": 4.5,
                    "skill_assessment_scores": {"Python": 90, "AWS": 85},
                    "connection_count": 120
                }
            },
            {
                "candidate_id": "C-1002",
                "profile": {
                    "anonymized_name": "Candidate B",
                    "headline": "Full Stack Engineer",
                    "summary": "Full Stack developer specializing in dot net ecosystems.",
                    "location": "Mumbai",
                    "country": "India",
                    "years_of_experience": 3.0,
                    "current_title": "Software Engineer",
                    "current_company": "Enterprise Solutions",
                    "current_company_size": "500-1000",
                    "current_industry": "Finance"
                },
                "career_history": [
                    {
                        "company": "Enterprise Solutions",
                        "title": "Software Engineer",
                        "start_date": "2023-06",
                        "end_date": "present",
                        "duration_months": 36,
                        "is_current": True,
                        "industry": "Finance",
                        "company_size": "500-1000",
                        "description": "Building web apps using ASP.NET Core and React."
                    }
                ],
                "skills": ["C#", "dot net", "SQL", "React", "college LAN"],
                "redrob_signals": {
                    "profile_completeness_score": 0.80,
                    "signup_date": "2023-05-01",
                    "last_active_date": "2026-05-30",
                    "open_to_work_flag": False,
                    "profile_views_received_30d": 12,
                    "applications_submitted_30d": 1,
                    "recruiter_response_rate": np.nan,
                    "avg_response_time_hours": np.nan,
                    "skill_assessment_scores": [75, 80],
                    "connection_count": 45
                }
            }
        ]
        with open(json_test_file, 'w', encoding='utf-8') as f:
            json.dump(synthetic_candidates, f, indent=2)
            
    # Generate mock DOCX if not present
    if not os.path.exists(docx_test_file):
        print(f"Generating synthetic job description in '{docx_test_file}'...")
        doc = docx.Document()
        doc.add_heading("Senior Backend Engineer (Python)", 0)
        doc.add_paragraph("We are looking for a Senior Backend Engineer to join our core AI and data platform team.")
        doc.add_paragraph("Requirements:")
        doc.add_paragraph("- Deep experience in Python programming and REST APIs development.")
        doc.add_paragraph("- Cloud experience with AWS and microservices architecture.")
        doc.add_paragraph("- Knowledge of containers (Docker/Kubernetes).")
        doc.save(docx_test_file)

    # Initialize and execute pipeline
    pipeline = TalentDataPipeline()
    print("\n--- Executing upgrade pipeline run ---")
    clean_profiles, clean_jobs = pipeline.run(json_test_file, docx_test_file)
    
    print("\n=== Validation Metrics ===")
    print(f"Clean Profiles Shape: {clean_profiles.shape}")
    print(f"Clean Jobs Shape:     {clean_jobs.shape}")
    print("\nMissing Value Count in Clean Profiles:")
    print(clean_profiles.isnull().sum())
    
    print("\nFeature Engineering Verification:")
    engineered = pipeline.engineer_features(clean_profiles)
    print(f"Engineered Shape  : {engineered.shape}")
    print(f"Engineered Columns: {list(engineered.columns)}")
    print(f"Anonymised (Dropped): {sorted(set(clean_profiles.columns) - set(engineered.columns))}")
    print("\nSkills Translation Verification:")
    print(engineered[['skills_raw', 'skills_clean']].to_string(index=False))
    print("\nDate-Based Career Trajectory Metrics:")
    traj_cols = ['candidate_id', 'total_duration_months', 'num_transitions', 'avg_tenure_months', 'career_velocity_score']
    traj_present = [c for c in traj_cols if c in engineered.columns]
    print(engineered[traj_present].to_string(index=False))
    print("\nNormalized Text Blob (first 200 chars):")
    if 'text_normalized' in engineered.columns:
        for _, row in engineered.iterrows():
            print(f"  {row['candidate_id']}: {str(row['text_normalized'])[:200]}")
