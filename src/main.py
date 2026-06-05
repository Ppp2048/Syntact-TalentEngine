"""
main.py
-------
Central orchestrator for the Syntact TalentEngine production pipeline.

Enforces strict offline mode (no Hugging Face network requests) and
streams the full hackathon dataset (data/candidates.jsonl) while
writing exactly 100 ranked rows to outputs/ranked_shortlist.csv.

Run from the project root:
    python src/main.py
"""

# ---------------------------------------------------------------------------
# OFFLINE COMPLIANCE — must be set before any transformers / HF import
# ---------------------------------------------------------------------------
import os
os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["HF_DATASETS_OFFLINE"] = "1"
os.environ["HF_HUB_DISABLE_SYMLINKS_FILTER"] = "1"

import sys
import logging
import time
from typing import List

import numpy as np
import pandas as pd

# ---------------------------------------------------------------------------
# Path bootstrap — project root and src/ both on sys.path so plain
# "import data_pipeline" resolves correctly (no "src." prefix needed).
# ---------------------------------------------------------------------------
_SRC_DIR  = os.path.dirname(os.path.abspath(__file__))
_PROJ_DIR = os.path.dirname(_SRC_DIR)

for _p in [_PROJ_DIR, _SRC_DIR]:
    if _p not in sys.path:
        sys.path.insert(0, _p)

# Local module imports (plain names — src/ is on sys.path)
from data_pipeline import TalentDataPipeline
from embedder import DualSignalEncoder
from ranker import PredictiveRankerEngine

# ---------------------------------------------------------------------------
# Logger
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s.%(msecs)03d  [%(levelname)-8s]  %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger("orchestrator")

# ---------------------------------------------------------------------------
# Path constants
# ---------------------------------------------------------------------------
DATA_DIR        = os.path.join(_PROJ_DIR, "data")
OUTPUTS_DIR     = os.path.join(_PROJ_DIR, "outputs")

# Production dataset — full 100k-candidate JSONL stream
CANDIDATES_PATH = os.path.join(DATA_DIR, "candidates.jsonl")
# Fallback to sample JSON if production file absent (dev / CI mode)
_SAMPLE_PATH    = os.path.join(DATA_DIR, "sample_candidates.json")

JOB_DESC_PATH   = os.path.join(DATA_DIR, "job_description.docx")
OUTPUT_CSV      = os.path.join(OUTPUTS_DIR, "ranked_shortlist.csv")

# Pipeline hyper-parameters
ENCODER_MODEL   = "all-MiniLM-L6-v2"
ENCODER_BATCH   = 32
TOP_K           = 100          # exactly 100 rows in the final CSV
RANDOM_STATE    = 42
# LightGBM hard limit: no single query group may exceed this row count.
# For 100k candidates we create ceil(100000 / 10000) = 10 equal groups.
_LGBM_MAX_GROUP = 10_000
# ---------------------------------------------------------------------------
# TRIAL MODE — set TRIAL_ROWS > 0 to run a fast verification subset.
# Set to 0 to process the full production dataset (100k candidates).
# When active, embeddings are cached separately as trial_embeddings.npy
# so the production cache (capability_embeddings.npy) is never polluted.
# ---------------------------------------------------------------------------
TRIAL_ROWS = 0   # PRODUCTION MODE: full 100k run. Set > 0 (e.g. 5_000) for a trial subset.


# ---------------------------------------------------------------------------
# DOCX reader
# ---------------------------------------------------------------------------

def read_docx(file_path: str) -> str:
    """Extracts plain text from a .docx job description file."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Job description not found: {file_path}")
    try:
        import docx as _docx
    except ImportError:
        raise ImportError(
            "python-docx is required. Install with: pip install python-docx"
        )
    doc = _docx.Document(file_path)
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(p.strip() for p in parts if p.strip())


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def main() -> None:
    t_start = time.perf_counter()
    os.makedirs(OUTPUTS_DIR, exist_ok=True)

    # -----------------------------------------------------------------------
    # Resolve candidates file (production JSONL  →  fallback JSON sample)
    # -----------------------------------------------------------------------
    if os.path.exists(CANDIDATES_PATH):
        candidates_path = CANDIDATES_PATH
        logger.info("Production dataset  : %s", candidates_path)
    elif os.path.exists(_SAMPLE_PATH):
        candidates_path = _SAMPLE_PATH
        logger.warning(
            "Production JSONL not found — falling back to sample: %s",
            _SAMPLE_PATH,
        )
    else:
        raise FileNotFoundError(
            f"No candidate data found. Expected:\n"
            f"  {CANDIDATES_PATH}\n  {_SAMPLE_PATH}"
        )

    # -----------------------------------------------------------------------
    # Stage 1 — Job description
    # -----------------------------------------------------------------------
    logger.info("=== Stage 1: Job Description Parsing ===")
    job_text = read_docx(JOB_DESC_PATH)
    logger.info("Job description parsed (%d chars).", len(job_text))

    # -----------------------------------------------------------------------
    # Stage 2 — Data ingestion, cleaning, feature engineering
    #           Streams JSONL line-by-line; safe for 100k+ records.
    # -----------------------------------------------------------------------
    logger.info("=== Stage 2: Data Ingestion & Feature Engineering ===")
    pipeline = TalentDataPipeline()

    logger.info("Streaming candidates from: %s", candidates_path)
    # Pass max_rows into the generator so the JSONL stream stops immediately
    # after TRIAL_ROWS records — no time wasted parsing the remaining 95k rows.
    _ingest_limit = TRIAL_ROWS if (TRIAL_ROWS and TRIAL_ROWS > 0) else None
    df_raw  = pipeline.load_and_clean_candidates(candidates_path, max_rows=_ingest_limit)
    logger.info("Loaded %d candidates after honeypot sweep.", len(df_raw))

    df_feat = pipeline.engineer_features(df_raw)
    logger.info("Feature engineering complete — shape: %s", df_feat.shape)

    # Ensure behavioural proxy columns exist
    for col, default in [
        ("profile_update_frequency",   0.0),
        ("submission_timestamp_index", 0.0),
        ("interaction_history_index",  0.0),
    ]:
        if col not in df_feat.columns:
            df_feat[col] = default
            logger.warning("Synthesised missing column: %s", col)

    # -----------------------------------------------------------------------
    # Mode announcement + cache filename routing
    # -----------------------------------------------------------------------
    if TRIAL_ROWS and TRIAL_ROWS > 0:
        cache_filename = "trial_embeddings.npy"
        logger.warning(
            "[TRIAL MODE] Running on first %d candidates (ingested via early-exit stream). "
            "Embeddings cached as '%s'. Set TRIAL_ROWS = 0 to restore full production run.",
            len(df_feat), cache_filename,
        )
    else:
        cache_filename = "capability_embeddings.npy"
        logger.info("[PRODUCTION MODE] Processing all %d candidates.", len(df_feat))

    # -----------------------------------------------------------------------
    # Stage 3 — Dual-signal encoding
    # -----------------------------------------------------------------------
    logger.info("=== Stage 3: Dual-Signal Encoding ===")

    # Pick richest available text column
    for _col in ("text_normalized", "skills_clean", "experience_text", "skills_raw"):
        if _col in df_feat.columns:
            text_col = _col
            break
    else:
        raise RuntimeError("No usable text column found in engineered profiles.")

    logger.info("Using text column: '%s'", text_col)
    candidate_texts: List[str] = (
        df_feat[text_col].fillna("").astype(str).tolist()
    )

    encoder = DualSignalEncoder(
        model_name=ENCODER_MODEL,
        batch_size=ENCODER_BATCH,
        normalize_embeddings=True,
        cache_filename=cache_filename,
    )

    composite           = encoder.build_composite_matrix(candidate_texts, df_feat)
    candidate_embeddings: np.ndarray = composite["capability"]   # (N, 384)
    intent_matrix:        np.ndarray = composite["intent"]       # (N, K)

    # Encode job description separately — is_job_desc=True bypasses the
    # candidate embedding cache so the production .npy file is never touched.
    job_embedding: np.ndarray = encoder.encode_capability(
        [job_text], is_job_desc=True
    )[0]

    logger.info(
        "Encoding complete — capability: %s | intent: %s | job: %s",
        candidate_embeddings.shape,
        intent_matrix.shape,
        job_embedding.shape,
    )

    # -----------------------------------------------------------------------
    # Stage 4 — LightGBM LambdaMART ranking
    # -----------------------------------------------------------------------
    logger.info("=== Stage 4: LambdaMART Ranking ===")

    ranker = PredictiveRankerEngine(top_k=TOP_K, random_state=RANDOM_STATE)

    # Cosine similarity between job embedding and every candidate
    cosine_sims   = ranker.compute_cosine_similarity(job_embedding, candidate_embeddings)
    feature_matrix = ranker.build_feature_matrix(cosine_sims, intent_matrix)

    # Pseudo-relevance labels: quantile-binned weighted combination
    intent_scores   = intent_matrix.mean(axis=1)
    combined_signal = 0.7 * cosine_sims + 0.3 * intent_scores
    q25, q50, q75  = np.quantile(combined_signal, [0.25, 0.50, 0.75])
    relevance_labels = np.digitize(
        combined_signal, [q25, q50, q75]
    ).astype(np.int32)

    # Fit LGBMRanker (LambdaMART)
    # ----------------------------------------------------------------
    # LightGBM raises an internal error if any single query group
    # exceeds 10,000 rows.  Split N candidates into sequential blocks
    # of _LGBM_MAX_GROUP rows each.  The last block absorbs any
    # remainder so that sum(group_sizes) == N exactly.
    # ----------------------------------------------------------------
    import math as _math
    n_total    = len(df_feat)
    n_groups   = _math.ceil(n_total / _LGBM_MAX_GROUP)
    # Build equal-size groups; remainder goes into the last group
    base_size, remainder = divmod(n_total, n_groups)
    group_sizes = np.array(
        [base_size + (1 if i < remainder else 0) for i in range(n_groups)],
        dtype=np.int32,
    )
    logger.info(
        "LGBMRanker group segmentation: %d candidates → %d groups "
        "(sizes: min=%d, max=%d, sum=%d)",
        n_total, n_groups,
        int(group_sizes.min()), int(group_sizes.max()), int(group_sizes.sum()),
    )
    assert group_sizes.sum() == n_total, "Group size mismatch — groups do not cover all candidates."
    ranker.fit(feature_matrix, relevance_labels, group_sizes)

    # Predict alignment scores for every candidate
    alignment_scores = ranker.score_candidates(feature_matrix)

    # -----------------------------------------------------------------------
    # Stage 5 — Shortlist generation with deterministic tie-breaking
    # -----------------------------------------------------------------------
    logger.info("=== Stage 5: Shortlist Generation ===")

    candidate_ids = df_feat["candidate_id"].tolist()
    shortlist = ranker.generate_shortlist(
        candidate_ids, cosine_sims, intent_matrix, alignment_scores
    )

    # NDCG@10 evaluation
    ndcg = ranker.evaluate_ndcg(relevance_labels, alignment_scores, k=10)
    logger.info("NDCG@10 = %.4f", ndcg)

    # -----------------------------------------------------------------------
    # Stage 6 — Persist exactly TOP_K rows to outputs/ranked_shortlist.csv
    # -----------------------------------------------------------------------
    logger.info("=== Stage 6: Persisting Submission CSV ===")
    saved_path = ranker.save_shortlist(
        shortlist, top_k=TOP_K, output_path=OUTPUT_CSV
    )

    # Verify output integrity
    written = pd.read_csv(saved_path)
    assert len(written) == min(TOP_K, len(shortlist)), (
        f"Expected {min(TOP_K, len(shortlist))} rows, got {len(written)}"
    )
    assert list(written.columns) == ["candidate_id", "rank", "score", "reasoning"], (
        f"Column mismatch: {list(written.columns)}"
    )

    elapsed = time.perf_counter() - t_start
    logger.info("=" * 65)
    logger.info("  PIPELINE COMPLETE")
    logger.info("  Candidates processed : %d", len(df_feat))
    logger.info("  Shortlist rows saved : %d", len(written))
    logger.info("  Output CSV           : %s", saved_path)
    logger.info("  NDCG@10              : %.4f", ndcg)
    logger.info("  Wall time            : %.2f s", elapsed)
    logger.info("=" * 65)

    # Top-5 preview
    logger.info("Top-5 ranked candidates:")
    for _, row in written.head(5).iterrows():
        logger.info(
            "  #%-3d  %-20s  score=%.4f",
            int(row["rank"]),
            str(row["candidate_id"]),
            float(row["score"]),
        )


if __name__ == "__main__":
    main()
